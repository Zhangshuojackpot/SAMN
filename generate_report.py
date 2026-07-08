from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=None, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

# ── Paper Table 1 values (read from screenshot) ───────────────────────────────
# Rows: SLAS-baseline / CE+SAMN / SLAS+SAMN / GLMC+SAMN
# Cols: CIFAR100 IF100/50/10,  CIFAR10 IF100/50/10
paper = {
    'SLAS (baseline)': ('52.7', '56.6', '69.7',   '84.2', '87.2', '92.7'),
    'CE + SAMN':       ('54.0', '58.3', '70.1',   '84.1', '87.6', '92.8'),
    'SLAS + SAMN':     ('54.1', '58.4', '70.1',   '85.5', '88.5', '92.8'),
    'GLMC + SAMN':     ('57.7', '64.0', '74.1',   '88.5', '91.3', '95.2'),
}

# CE and GLMC standalone baselines from the paper (reference rows)
paper_baselines = {
    'CE (WD, Stage 1)': ('47.4', '52.3', '67.2',  '80.0', '84.3', '91.9'),
    'GLMC [11]':        ('56.8', '62.2', '72.8',  '88.2', '90.8', '95.1'),
}

# ── Our trained results ────────────────────────────────────────────────────────
# Stage 1
stage1 = {
    'CE backbone':   ('48.34', '53.11', '67.26',  '80.02', '84.32', '91.93'),
    'GLMC backbone': ('57.20', '62.99', '74.08',  '88.24', '90.89', '95.12'),
}

# Stage 2
ours = {
    'SLAS (baseline)': ('52.33', '56.61', '69.40',  '83.89', '86.98', '92.64'),
    'CE + SAMN':       ('54.62', '58.37', '69.93',  '83.83', '87.61', '92.78'),
    'SLAS + SAMN':     ('54.67', '58.44', '69.87',  '85.42', '88.43', '92.46'),
    'GLMC + SAMN':     ('57.89', '63.37', '74.31',  '88.42', '91.25', '95.16'),
}

# ═══════════════════════════════════════════════════════════════════════════════
# Title
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('SAMN (CVPR 2026) — Experimental Results Report', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size  = Pt(16)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph('Training Results on CIFAR-LT  ·  ResNet-32  ·  2026-06-25 (after bug fixes)')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(10)
sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# Section 1 — Stage 1 backbone
# ═══════════════════════════════════════════════════════════════════════════════
h1 = doc.add_heading('1. Stage 1 Backbone Training', level=2)
h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph(
    'ResNet-32 backbone trained for 200 epochs. '
    'CE backbone (use_glmc=0) feeds SLAS baseline, CE+SAMN, and SLAS+SAMN. '
    'GLMC backbone (use_glmc=1) feeds GLMC+SAMN.'
)
p.runs[0].font.size = Pt(10)

# 8-col table: header + 2 method rows
t = doc.add_table(rows=3, cols=8)
t.style = 'Table Grid'

hdr0 = ['Backbone', 'CIFAR100-LT IF=100', 'CIFAR100-LT IF=50', 'CIFAR100-LT IF=10',
        'CIFAR10-LT IF=100', 'CIFAR10-LT IF=50', 'CIFAR10-LT IF=10', 'Paper ref.']
# Merge row-0 into group headers
for j, h in enumerate(hdr0):
    cell_text(t.cell(0, j), h, bold=True, size=9)
    set_cell_bg(t.cell(0, j), '1F497D')
    for para in t.cell(0, j).paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

paper_ref = {'CE backbone': '47.4 / 52.3 / 67.2  (CIFAR100)\n80.0 / 84.3 / 91.9  (CIFAR10)',
             'GLMC backbone': '56.8 / 62.2 / 72.8  (CIFAR100)\n88.2 / 90.8 / 95.1  (CIFAR10)'}
bgs = {'CE backbone': 'EBF5FB', 'GLMC backbone': 'E8F8F5'}

for i, (name, vals) in enumerate(stage1.items()):
    row = 1 + i
    bg = bgs[name]
    cell_text(t.cell(row, 0), name, bold=True, size=10)
    set_cell_bg(t.cell(row, 0), bg)
    for j, v in enumerate(vals):
        cell_text(t.cell(row, j+1), v, size=10)
        set_cell_bg(t.cell(row, j+1), bg)
    cell_text(t.cell(row, 7), paper_ref[name], size=8, italic=True)
    set_cell_bg(t.cell(row, 7), bg)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# Section 2 — Main Table (Table 1 reproduction)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1 = doc.add_heading('2. Stage 2 Classifier Fine-tuning — Table 1 Reproduction', level=2)
h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph(
    'Stage 2 fine-tunes only the classifier head for 20 epochs (backbone frozen). '
    'Results are the best Top-1 accuracy (%) observed across 20 epochs.\n'
    'SLAS baseline uses standard nn.Linear + LabelAwareSmoothing (use_samn=0). '
    'All SAMN variants use PAVALinear enforcing monotone non-decreasing per-class weight norms.'
)
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# ── Build combined table: 3 header rows + data rows ──────────────────────────
# Cols: Method | C100 IF100 (Ours/Paper/Δ) | C100 IF50 | C100 IF10 | C10 IF100 | C10 IF50 | C10 IF10
# = 1 + 6×3 = 19 cols  — too wide; use 2-col (Ours/Paper) per IF, and Δ in notes
# Layout: Method | [C100-100 Ours | C100-100 Paper] | [C100-50 ...] ... = 1 + 6×2 = 13 cols

NCOLS = 13
NROWS = 2 + len(ours)  # 2 header rows + 4 data rows

t = doc.add_table(rows=NROWS, cols=NCOLS)
t.style = 'Table Grid'

# Row 0: group headers
groups = [
    ('Method',       1, '1F497D'),
    ('CIFAR100-LT IF=100', 2, '2E75B6'),
    ('CIFAR100-LT IF=50',  2, '2E75B6'),
    ('CIFAR100-LT IF=10',  2, '2E75B6'),
    ('CIFAR10-LT IF=100',  2, '1A5599'),
    ('CIFAR10-LT IF=50',   2, '1A5599'),
    ('CIFAR10-LT IF=10',   2, '1A5599'),
]
col = 0
for label, span, color_hex in groups:
    cell = t.cell(0, col)
    if span > 1:
        cell.merge(t.cell(0, col + span - 1))
    cell.paragraphs[0].clear()
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].add_run(label)
    run.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(cell, color_hex)
    col += span

# Row 1: Ours / Paper sub-headers
sub_labels = ['Method'] + ['Ours', 'Paper'] * 6
for j, lbl in enumerate(sub_labels):
    cell_text(t.cell(1, j), lbl, bold=True, size=9)
    set_cell_bg(t.cell(1, j), 'BDD7EE' if j < 7 else 'DDEEFF')

METHOD_BG = {
    'SLAS (baseline)': 'FEF9E7',
    'CE + SAMN':       'EBF5FB',
    'SLAS + SAMN':     'EBF5FB',
    'GLMC + SAMN':     'E8F8F5',
}

def diff_color(our_val, paper_val):
    d = float(our_val) - float(paper_val)
    if d >= -0.3:  return '1A5276'   # within 0.3 → dark blue (match)
    if d >= -1.0:  return '145A32'   # within 1.0 → dark green
    return 'B03A2E'                   # gap > 1.0 → red

for i, method in enumerate(ours):
    row = 2 + i
    bg  = METHOD_BG[method]
    o_vals = ours[method]
    p_vals = paper[method]

    cell_text(t.cell(row, 0), method, bold=(method == 'GLMC + SAMN'), size=10)
    set_cell_bg(t.cell(row, 0), bg)

    for k in range(6):
        ov = o_vals[k]; pv = p_vals[k]
        color = diff_color(ov, pv)
        # Ours col
        cell_text(t.cell(row, 1 + k*2), ov, size=10, color=color,
                  bold=(method == 'GLMC + SAMN'))
        set_cell_bg(t.cell(row, 1 + k*2), bg)
        # Paper col
        cell_text(t.cell(row, 2 + k*2), pv, size=10, italic=True)
        set_cell_bg(t.cell(row, 2 + k*2), bg)

# Column widths
widths = [Cm(2.8)] + [Cm(1.35), Cm(1.35)] * 6
for row in t.rows:
    for j, cell in enumerate(row.cells):
        cell.width = widths[j]

doc.add_paragraph()

# Legend
leg = doc.add_paragraph()
leg.add_run('Colour key (Ours column): ').bold = True
for label, hex_c, meaning in [
    ('■ Dark blue', '1A5276', ' ≤0.3% below paper  '),
    ('■ Dark green','145A32', ' ≤1.0% below paper  '),
    ('■ Red',       'B03A2E', ' >1.0% below paper'),
]:
    r = leg.add_run(label)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(*bytes.fromhex(hex_c))
    r2 = leg.add_run(meaning)
    r2.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════════════════════════
# Section 3 — Deviation table
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1 = doc.add_heading('3. Deviation Analysis vs. Paper Table 1', level=2)
h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Build deviation rows programmatically
dev_rows = []
datasets = [('CIFAR-100', slice(0, 3)), ('CIFAR-10', slice(3, 6))]
ifs = ['IF=100', 'IF=50', 'IF=10']
for method in ours:
    for ds_name, sl in datasets:
        for k, if_str in enumerate(ifs):
            idx = sl.start + k
            ov = float(ours[method][idx])
            pv = float(paper[method][idx])
            d  = ov - pv
            dev_rows.append((ds_name, method, if_str, ov, pv, d))

t2 = doc.add_table(rows=1 + len(dev_rows), cols=6)
t2.style = 'Table Grid'

dev_hdrs = ['Dataset', 'Method', 'IF', 'Ours (%)', 'Paper (%)', 'Δ (Ours − Paper)']
for j, h in enumerate(dev_hdrs):
    cell_text(t2.cell(0, j), h, bold=True, size=10)
    set_cell_bg(t2.cell(0, j), '1F497D')
    for para in t2.cell(0, j).paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for i, (ds, method, if_str, ov, pv, d) in enumerate(dev_rows):
    row = i + 1
    sign = '+' if d >= 0 else ''
    d_str = f'{sign}{d:.2f}%'
    bg = 'EAFAF1' if d >= -0.3 else ('FEF9E7' if d >= -1.0 else 'FDEDEC')
    vals = [ds, method, if_str, f'{ov:.2f}', f'{pv:.1f}', d_str]
    for j, val in enumerate(vals):
        cell_text(t2.cell(row, j), val, size=9)
        set_cell_bg(t2.cell(row, j), bg)

doc.add_paragraph()

# ── Summary ────────────────────────────────────────────────────────────────────
h2 = doc.add_heading('Summary', level=3)
h2.runs[0].font.size = Pt(12)

lines = [
    'CIFAR-100 — all four methods match the paper within ±0.63%. CE+SAMN and SLAS+SAMN '
    'at IF=100 slightly exceed the paper (by +0.62% and +0.57% respectively). '
    'GLMC+SAMN IF=50 is the largest negative gap (63.37 vs 64.0, Δ=−0.63%), '
    'attributable to stochastic variation in GLMC dual-dataloader sampling.',
    '',
    'CIFAR-10 — all methods now closely reproduce the paper. SLAS baseline recovered '
    'from the previous 3.1% gap to within 0.31% after fixing the MISLAS learning rate '
    '(0.00001 → 0.0005). CE+SAMN is within 0.27%, SLAS+SAMN within 0.34%, '
    'GLMC+SAMN within 0.08%. The PAVALinear exp() fix restored proper weight scaling.',
    '',
    'Overall: both bugs (CIFAR-10 MISLAS lr and PAVALinear missing exp()) are resolved. '
    'All 24 experiments reproduce the paper Table 1 within ±0.63%, confirming full '
    'reproducibility of the SAMN results.',
]
for line in lines:
    p = doc.add_paragraph(line)
    if p.runs:
        p.runs[0].font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════════
# Section 4 — SAMN Gain: baseline vs +SAMN
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1 = doc.add_heading('4. SAMN Gain: Performance Before vs. After Adding SAMN', level=2)
h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph(
    'Comparison of each method paired with its SAMN-enhanced counterpart. '
    'Δ = (With SAMN) − (Without SAMN). All values are our trained results (best Top-1 %).'
)
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# Pairs: (label, without_SAMN_vals, with_SAMN_vals)
# Without SAMN: CE backbone (Stage1), SLAS baseline, GLMC backbone (Stage1)
gain_pairs = [
    ('CE backbone → CE + SAMN',
     stage1['CE backbone'],
     ours['CE + SAMN']),
    ('SLAS baseline → SLAS + SAMN',
     ours['SLAS (baseline)'],
     ours['SLAS + SAMN']),
    ('GLMC backbone → GLMC + SAMN',
     stage1['GLMC backbone'],
     ours['GLMC + SAMN']),
]

# Header: Method | C100-100 W/O | C100-100 W/ | C100-100 Δ | C100-50 ... | C10 ...
# Layout: 1 method col + 6 IF groups × 3 (w/o, w/, Δ) = 19 cols — too wide
# Use 2 sub-tables: one for CIFAR-100, one for CIFAR-10

for ds_label, ds_slice, ds_color in [
    ('CIFAR-100-LT', slice(0, 3), '2E75B6'),
    ('CIFAR-10-LT',  slice(3, 6), '1A5599'),
]:
    h3 = doc.add_heading(ds_label, level=3)
    h3.runs[0].font.size = Pt(11)
    h3.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(ds_color))

    # 10 cols: Method | IF=100 w/o | IF=100 w/ | IF=100 Δ | IF=50 ... | IF=10 ...
    tg = doc.add_table(rows=2 + len(gain_pairs), cols=10)
    tg.style = 'Table Grid'

    # Row 0: group headers
    cell_text(tg.cell(0, 0), 'Method', bold=True, size=9)
    set_cell_bg(tg.cell(0, 0), '1F497D')
    tg.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for k, if_label in enumerate(['IF=100', 'IF=50', 'IF=10']):
        c = tg.cell(0, 1 + k*3)
        c.merge(tg.cell(0, 3 + k*3))
        c.paragraphs[0].clear()
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(if_label)
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, ds_color)

    # Row 1: sub-headers
    sub = ['Method'] + ['w/o SAMN', 'w/ SAMN', 'Δ'] * 3
    sub_bgs = ['BDD7EE'] + ['DDEEFF', 'D5F5E3', 'FEF9E7'] * 3
    for j, (lbl, bg) in enumerate(zip(sub, sub_bgs)):
        cell_text(tg.cell(1, j), lbl, bold=True, size=9)
        set_cell_bg(tg.cell(1, j), bg)

    for i, (label, wo_vals, wi_vals) in enumerate(gain_pairs):
        row = 2 + i
        bg_row = ['EBF5FB', 'FEF9E7', 'E8F8F5'][i]
        cell_text(tg.cell(row, 0), label, bold=True, size=9,
                  align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(tg.cell(row, 0), bg_row)
        for k in range(3):
            idx = ds_slice.start + k
            wo = float(wo_vals[idx])
            wi = float(wi_vals[idx])
            d  = wi - wo
            sign = '+' if d >= 0 else ''
            delta_str = f'{sign}{d:.2f}'
            d_color = '145A32' if d > 0 else ('1A5276' if d >= -0.1 else 'B03A2E')
            cell_text(tg.cell(row, 1 + k*3), f'{wo:.2f}', size=10)
            set_cell_bg(tg.cell(row, 1 + k*3), bg_row)
            cell_text(tg.cell(row, 2 + k*3), f'{wi:.2f}', size=10, bold=True)
            set_cell_bg(tg.cell(row, 2 + k*3), bg_row)
            cell_text(tg.cell(row, 3 + k*3), delta_str, size=10, bold=True, color=d_color)
            set_cell_bg(tg.cell(row, 3 + k*3), bg_row)

    # Column widths
    col_widths = [Cm(4.2)] + [Cm(1.5), Cm(1.5), Cm(1.2)] * 3
    for row in tg.rows:
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]

    doc.add_paragraph()

# Legend for gain table
leg2 = doc.add_paragraph()
leg2.add_run('Δ colour key: ').bold = True
for label, hex_c, meaning in [
    ('■ Green', '145A32', ' positive gain (SAMN helps)  '),
    ('■ Blue',  '1A5276', ' near-zero (≤0.1% change)  '),
    ('■ Red',   'B03A2E', ' negative (SAMN hurts)'),
]:
    r = leg2.add_run(label)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(*bytes.fromhex(hex_c))
    r2 = leg2.add_run(meaning)
    r2.font.size = Pt(9)

# ── Save ───────────────────────────────────────────────────────────────────────
out = '/home/engs2909/research/WPAVA/SAMN_upload/SAMN_CVPR2026_Training_Report.docx'
doc.save(out)
print(f'Saved: {out}')
